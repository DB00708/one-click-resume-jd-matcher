from docx import Document


def create_resume(resume_based_on_updated_skills, path):
    doc = Document()

    doc.add_heading(resume_based_on_updated_skills.get('full_name'), 0)

    if 'user_email' in resume_based_on_updated_skills or 'user_mobile_no' in resume_based_on_updated_skills:
        doc.add_heading('Contact Information', level=1)
        if 'user_email' in resume_based_on_updated_skills:
            doc.add_paragraph(f"Email: {resume_based_on_updated_skills['user_email']}")
        if 'user_mobile_no' in resume_based_on_updated_skills:
            doc.add_paragraph(f"Mobile: {resume_based_on_updated_skills['user_mobile_no']}")

    if 'skills' in resume_based_on_updated_skills:
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(', '.join(resume_based_on_updated_skills['skills']))

    if 'education' in resume_based_on_updated_skills:
        doc.add_heading('Education', level=1)
        for education in resume_based_on_updated_skills['education']:
            doc.add_heading(education.get('degree', 'No Degree Provided'), level=2)
            if 'university' in education:
                doc.add_paragraph(f"University: {education['university']}")
            if 'board' in education:
                doc.add_paragraph(f"Board: {education['board']}")
            if 'school' in education:
                doc.add_paragraph(f"School: {education['school']}")
            if 'duration' in education:
                doc.add_paragraph(f"Duration: {education['duration']}")
            if 'year' in education:
                doc.add_paragraph(f"Year: {education['year']}")
            if 'percentage' in education:
                doc.add_paragraph(f"Percentage: {education['percentage']}")
            if 'grade' in education:
                doc.add_paragraph(f"Grade: {education['grade']}")

    if 'projects' in resume_based_on_updated_skills:
        doc.add_heading('Projects', level=1)
        for project in resume_based_on_updated_skills['projects']:
            doc.add_heading(project.get('title', 'No Title Provided'), level=2)
            doc.add_paragraph(f"Duration: {project.get('duration', 'No Duration Provided')}")
            doc.add_paragraph(f"Description: {project.get('description', 'No Description Provided')}")

    if 'responsibility' in resume_based_on_updated_skills:
        doc.add_heading('Responsibilities', level=1)
        for responsibility in resume_based_on_updated_skills['responsibility']:
            doc.add_heading(responsibility.get('role', 'No Role Provided'), level=2)
            doc.add_paragraph(f"Organization: {responsibility.get('organization', 'No Organization Provided')}")
            doc.add_paragraph(f"Duration: {responsibility.get('duration', 'No Duration Provided')}")
            if 'responsibilities' in responsibility:
                doc.add_paragraph('Responsibilities:')
                for resp in responsibility['responsibilities']:
                    doc.add_paragraph(f"- {resp}")

    if 'achievements' in resume_based_on_updated_skills:
        doc.add_heading('Achievements', level=1)
        for achievement in resume_based_on_updated_skills['achievements']:
            doc.add_heading(achievement.get('title', 'No Title Provided'), level=2)
            doc.add_paragraph(f"Organization: {achievement.get('organization', 'No Organization Provided')}")
            doc.add_paragraph(f"Details: {achievement.get('details', 'No Details Provided')}")

    if 'training' in resume_based_on_updated_skills:
        doc.add_heading('Training', level=1)
        for training in resume_based_on_updated_skills['training']:
            doc.add_heading(training.get('title', 'No Title Provided'), level=2)
            doc.add_paragraph(f"Duration: {training.get('duration', 'No Duration Provided')}")
            doc.add_paragraph(f"Details: {training.get('details', 'No Details Provided')}")

    doc.save(path)
